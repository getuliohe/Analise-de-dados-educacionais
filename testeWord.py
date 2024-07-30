import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Step 1: Create and save a plot using matplotlib
plt.figure(figsize=(6, 4))
plt.plot([1, 2, 3, 4], [1, 4, 2, 3])
plt.title("Sample Plot")
plt.xlabel("X-axis")
plt.ylabel("Y-axis")
plot_path = "plot.png"
plt.savefig(plot_path)
plt.close()

# Step 2: Create a Word document using python-docx
doc = Document()

# Add a title
doc.add_heading("Document Title", 0)

# Add some text
doc.add_paragraph("This is a sample text added to the document.")

# Insert the plot image
doc.add_heading("Plot Example", level=1)
doc.add_paragraph("Below is an example of a plot generated using matplotlib:")
doc.add_picture(plot_path, width=Inches(6))

# Save the document
doc.save("sample_document.docx")

print("Document created successfully.")
