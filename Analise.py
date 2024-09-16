import matplotlib.pyplot as plt
import numpy as np
from DadosAnalisados import DadosAnalisados
from docx import Document
from docx.shared import Inches
import math
from io import BytesIO


class Analise:
    @staticmethod
    def analise(tabela1, tabela2, anoSemestre, nomeProcura):
        doc = Document()
        
        Analise.realizarAnaliseDeUmDado(doc, tabela1.media, tabela2.media, anoSemestre, nomeProcura, "media")
        Analise.realizarAnaliseDeUmDado(doc, tabela1.totalFaltas, tabela2.totalFaltas, anoSemestre, nomeProcura, "Total de faltas")
        Analise.gerar_documento(doc, nomeProcura, anoSemestre)
    
    @staticmethod
    def realizarAnaliseDeUmDado(doc, dado1, dado2, anoSemestre, nomeProcura, nomeDado):
        dado1Analisado = Analise.retornaListaAnalisada(dado1)

        dado2Analisado = Analise.retornaListaAnalisada(dado2)
        
        medias = [dado1Analisado.media,dado2Analisado.media]

        listaPathGrafico = []
        listaPathGrafico.append(Analise.gerarGraficoBoxPLot(dado1, dado2, anoSemestre, nomeDado,))
        listaPathGrafico.append(Analise.gerarGraficosBarra(medias, anoSemestre, nomeDado, nomeProcura))

        Analise.adicionar_paragrafo_documento(doc, dado1Analisado, dado2Analisado, anoSemestre, nomeProcura, listaPathGrafico)

    @staticmethod
    def retornaListaAnalisada(lista):
        listaAnalisada = DadosAnalisados(
                                    len(lista),
                                    Analise.calculaMedia(lista),
                                    Analise.calcularMediana(lista),
                                    Analise.calcularVariancia(lista),
                                    Analise.calcularDesvioPadrao(lista),
                                    Analise.calcularDesvioMedioAbsoluto(lista),
                                    Analise.calcularAmplitude(lista),
                                    Analise.calcularCoeficienteDeVaricao(lista),
                                    min(Analise.filtrarNulos(lista)),
                                    max(Analise.filtrarNulos(lista)),
                                    Analise.calcularQuadrantes(lista)[0],
                                    Analise.calcularQuadrantes(lista)[1],
                                    Analise.calcularQuadrantes(lista)[2],
                                    Analise.calcularTotal(lista)
                                )
        
        return listaAnalisada


    @staticmethod
    def calculaMedia(lista):
        media = 0
        listaFiltrada = Analise.filtrarNulos(lista)
        totalLinhas = len(listaFiltrada)

        for linha in listaFiltrada:
                media += linha

        media /= totalLinhas
        return media
    
    @staticmethod
    def calcularQuadrantes(lista):
        # Filtrar None e ordenar a lista
        listaEmRol = sorted(filter(None, lista))
        
        def calcularQ1(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[n//4 - 1] + lista[n//4]) / 2
            else:
                return lista[n//4]

        def calcularQ2(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[n//2 - 1] + lista[n//2]) / 2
            else:
                return lista[n//2]

        def calcularQ3(lista):
            n = len(lista)
            if n % 2 == 0:
                return (lista[(3*n)//4 - 1] + lista[(3*n)//4]) / 2
            else:
                return lista[(3*n)//4]
            
        return calcularQ1(lista), calcularQ2(lista), calcularQ3(lista)

    
    @staticmethod
    def calcularMediana(lista):
        listaEmRol = sorted(Analise.filtrarNulos(lista))
        n = len(listaEmRol)

        if n % 2 != 0:
            return listaEmRol[(n // 2)]
        else:
            return (listaEmRol[(n // 2) - 1] + listaEmRol[(n // 2)])/2
    
    @staticmethod
    def calcularVariancia(lista):
        listaFiltrada = Analise.filtrarNulos(lista)
        media = Analise.calculaMedia(lista)
        total = 0

        for item in listaFiltrada:
            total += (item - media)**2
        
        return total/(len(listaFiltrada)- 1)
    
    @staticmethod
    def calcularDesvioPadrao(lista):
        return math.sqrt(Analise.calcularVariancia(lista))
    
    @staticmethod
    def calcularDesvioMedioAbsoluto(lista):
        listaFiltrada = Analise.filtrarNulos(lista)
        media = Analise.calculaMedia(lista)
        total = [abs(x - media) for x in listaFiltrada]
        
        return sum(total)/len(listaFiltrada)
    
    @staticmethod
    def calcularAmplitude(lista):
        return (max(Analise.filtrarNulos(lista)) - min(Analise.filtrarNulos(lista)))
    
    @staticmethod
    def calcularCoeficienteDeVaricao(lista):
        media = Analise.calculaMedia(lista)

        return (Analise.calcularDesvioPadrao(lista)/ media) * 100
    
    @staticmethod
    def filtrarNulos(lista):
        listaSemNulos = []

        for item in lista:
            if item != None or item == 0:
                listaSemNulos.append(item)

        return listaSemNulos
    
    @staticmethod
    def calcularTotal(lista):
        total = 0

        for item in lista:
            if item:
                total += item

        return total
    
    def gerarGraficoBoxPLot(lista1, lista2, anoSemestre, dadoAnalisado):

        plt.figure(figsize=(8,6))
        plt.boxplot([Analise.filtrarNulos(lista1), Analise.filtrarNulos(lista2)], patch_artist=True, showmeans=True, showfliers=True)
        plt.title(f"Boxplox da {dadoAnalisado} {anoSemestre[0]}X{anoSemestre[1]}")
        plt.ylabel('notas')
        plt.xticks([1, 2], [f"{anoSemestre[0]}", f"{anoSemestre[1]}"])
        plt.grid(True)
        plt.yticks(range(11))

        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)


        return img_stream

    def gerarGraficosBarra(valores, anoSemestre, dadoanalisado, listaAnalisada):
        plt.bar(anoSemestre, valores)

        plt.xlabel('Ano/Semestre')
        plt.ylabel('Média')
        plt.title('Gráfico de Barras')

        plt.grid(True)
        plt.yticks(range(11))

        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)

        return img_stream
    
    def adicionar_paragrafo_documento(doc, lista1, lista2, anos, nomeProcura, streamsGraficos):
        
        doc.add_heading(f"Analise de {nomeProcura} {anos[0]} X {nomeProcura} {anos[1]}")
        doc.add_paragraph(f"\tDado analisado:{nomeProcura}\n\n" +
                            f"Calculo: resultado {anos[0]} X resultado {anos[1]}" +
                            f"n = {lista1.tamanhoDaLista} X {lista2.tamanhoDaLista}\n" +
                            f"média = {lista1.media} X {lista2.media}\n" +
                            f"mediana = {lista1.mediana} X {lista2.mediana}\n" +
                            f"mínimo = {lista1.minimo} X {lista2.minimo}\n" +
                            f"quartil 1 = {lista1.quartil1} X {lista2.quartil1}\n" +
                            f"quartil 2 = {lista1.quartil2} X {lista2.quartil2}\n" +
                            f"quartil 3 = {lista1.quartil3} X {lista2.quartil3}\n" +
                            f"máximo = {lista1.maximo} X {lista2.maximo}\n" +
                            f"variancia = {lista1.variancia} X {lista2.variancia}\n" +
                            f"desvio padrão = {lista1.desvioPadrao} X {lista2.tamanhoDaLista}\n" +
                            f"desvio médio absoluto = {lista1.desvioMedioAbsoluto} X {lista2.desvioMedioAbsoluto}\n" +
                            f"amplitude = {lista1.amplitude} X {lista2.amplitude}\n" +
                            f"total = {lista1.total} X {lista2.total}")
        
        for stream in streamsGraficos:
            doc.add_picture(stream, width=Inches(6))
            
    def gerar_documento(doc, nomeProcura, anoSemestre):
        doc.save(f"{nomeProcura + anoSemestre[0] + anoSemestre[1]}.docx")